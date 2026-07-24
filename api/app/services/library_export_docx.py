"""Export Library course lessons to a formatted DOCX (all courses)."""
from __future__ import annotations

import io
import re
from typing import BinaryIO

from app.services.library_catalog import LessonDetail, _is_publishable, _lesson_sort_key, get_lesson, list_lessons


def _require_docx():
    try:
        from docx import Document  # noqa: F401
        from docx.shared import Pt, Inches  # noqa: F401
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is required for DOCX export. Install: pip install python-docx"
        ) from exc


def _strip_md_inline(text: str) -> list[tuple[str, bool, bool]]:
    if not text:
        return []
    pattern = re.compile(
        r"(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|__(.+?)__|\*(.+?)\*|_(.+?)_|`(.+?)`)"
    )
    out: list[tuple[str, bool, bool]] = []
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            out.append((text[pos : m.start()], False, False))
        if m.group(2) is not None:
            out.append((m.group(2), True, True))
        elif m.group(3) is not None:
            out.append((m.group(3), True, False))
        elif m.group(4) is not None:
            out.append((m.group(4), True, False))
        elif m.group(5) is not None:
            out.append((m.group(5), False, True))
        elif m.group(6) is not None:
            out.append((m.group(6), False, True))
        elif m.group(7) is not None:
            out.append((m.group(7), False, False))
        pos = m.end()
    if pos < len(text):
        out.append((text[pos:], False, False))
    return out or [(text, False, False)]


def _add_runs(paragraph, text: str) -> None:
    for chunk, bold, italic in _strip_md_inline(text):
        if not chunk:
            continue
        run = paragraph.add_run(chunk)
        run.bold = bold
        run.italic = italic


def _add_picture(doc, src: str, alt: str = "") -> None:
    from docx.shared import Inches

    from app.services.library_media import resolve_library_file

    path = None
    if src.startswith("/api/v1/library/files/"):
        path = resolve_library_file(src[len("/api/v1/library/files/") :])
    if path and path.is_file():
        try:
            doc.add_picture(str(path), width=Inches(5.5))
            if alt:
                p = doc.add_paragraph()
                run = p.add_run(alt)
                run.italic = True
            return
        except Exception:
            pass
    p = doc.add_paragraph()
    _add_runs(p, f"[Image: {alt or src}]")


def _add_md_table(doc, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(cols):
            cell = table.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            _add_runs(p, row[ci] if ci < len(row) else "")


def _add_body(doc, body: str) -> None:
    """Render markdown-ish lesson body into the document."""
    lines = (body or "").replace("\r\n", "\n").split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        img = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)$", stripped)
        if img:
            _add_picture(doc, img.group(2).strip(), alt=img.group(1).strip())
            i += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            table_lines: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            parsed: list[list[str]] = []
            for tl in table_lines:
                if re.match(r"^\|?[\s:|-]+\|?$", tl) and "-" in tl:
                    continue
                s = tl.strip()
                if s.startswith("|"):
                    s = s[1:]
                if s.endswith("|"):
                    s = s[:-1]
                parsed.append([c.strip() for c in s.split("|")])
            if parsed:
                _add_md_table(doc, parsed)
            continue

        hm = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if hm:
            level = min(len(hm.group(1)), 3)
            style = f"Heading {level + 1}"
            p = doc.add_paragraph(style=style)
            _add_runs(p, hm.group(2).strip())
            i += 1
            continue

        if re.match(r"^(-{3,}|\*{3,}|_{3,})$", stripped):
            doc.add_paragraph("─" * 24)
            i += 1
            continue

        um = re.match(r"^[-*+]\s+(.+)$", stripped)
        if um:
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, um.group(1).strip())
            i += 1
            continue

        om = re.match(r"^(\d+)[.)]\s+(.+)$", stripped)
        if om:
            p = doc.add_paragraph(style="List Number")
            _add_runs(p, om.group(2).strip())
            i += 1
            continue

        if stripped.startswith(">"):
            quote = re.sub(r"^>\s?", "", stripped)
            p = doc.add_paragraph()
            run = p.add_run(quote)
            run.italic = True
            i += 1
            continue

        buf = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt:
                break
            if re.match(r"^(#{1,6})\s+", nxt):
                break
            if re.match(r"^[-*+]\s+", nxt):
                break
            if re.match(r"^\d+[.)]\s+", nxt):
                break
            if nxt.startswith(">"):
                break
            if re.match(r"^(-{3,}|\*{3,}|_{3,})$", nxt):
                break
            if nxt.startswith("|"):
                break
            if re.match(r"^!\[", nxt):
                break
            buf.append(nxt)
            i += 1
        p = doc.add_paragraph()
        _add_runs(p, " ".join(buf))


def build_course_docx(course_id: str) -> tuple[bytes, str]:
    """
    Build a DOCX for all publishable lessons in a course.
    Returns (bytes, filename).
    """
    _require_docx()
    from docx import Document
    from docx.shared import Pt

    items, _, _ = list_lessons(course=course_id, published_only=True)
    if not items:
        raise ValueError(f"No publishable lessons for course {course_id!r}")

    lessons: list[LessonDetail] = []
    for summary in items:
        detail = get_lesson(summary.id)
        if detail and _is_publishable(detail):
            lessons.append(detail)
    lessons.sort(key=_lesson_sort_key)
    if not lessons:
        raise ValueError(f"No publishable lessons for course {course_id!r}")

    course_name = lessons[0].course or course_id
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading(course_name, level=0)
    title.runs[0].bold = True if title.runs else False

    current_module: str | None = None
    for lesson in lessons:
        module = (lesson.category or "").strip() or "General"
        if module != current_module:
            current_module = module
            doc.add_heading(module, level=1)
        kind_label = (lesson.kind or "text").capitalize()
        doc.add_heading(lesson.title or lesson.id, level=2)
        meta = doc.add_paragraph()
        run = meta.add_run(f"{kind_label}")
        run.italic = True
        run.font.size = Pt(9)
        body = (lesson.body or "").strip()
        if body:
            _add_body(doc, body)
        else:
            doc.add_paragraph("(No text content)")

    buf = io.BytesIO()
    doc.save(buf)
    data = buf.getvalue()
    safe = re.sub(r"[^\w\-]+", "-", course_name.strip().lower()).strip("-") or course_id
    filename = f"{safe}.docx"
    return data, filename


def write_course_docx(course_id: str, dest: BinaryIO) -> str:
    data, filename = build_course_docx(course_id)
    dest.write(data)
    return filename
